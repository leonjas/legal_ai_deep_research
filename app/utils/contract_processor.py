"""
Contract Processing Utilities
Handles contract file processing, text extraction, and formatting
"""

try:
    import PyMuPDF as fitz  # fitz
except ImportError:
    try:
        import fitz  # Alternative import
    except ImportError:
        fitz = None
import docx
import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import mimetypes
from io import BytesIO

class ContractProcessor:
    """
    Utility class for processing contract files and extracting text
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx,  # Will try docx method
            '.txt': self._extract_from_txt
        }
    
    def extract_text_from_file(self, file_path: str = None, file_content: bytes = None, 
                              filename: str = None) -> Tuple[str, Dict]:
        """
        Extract text from a contract file
        
        Args:
            file_path: Path to the file (if reading from disk)
            file_content: File content as bytes (if reading from memory)
            filename: Original filename for type detection
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            if file_path:
                file_extension = Path(file_path).suffix.lower()
                with open(file_path, 'rb') as f:
                    content = f.read()
            elif file_content and filename:
                file_extension = Path(filename).suffix.lower()
                content = file_content
            else:
                raise ValueError("Either file_path or (file_content + filename) must be provided")
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract text using appropriate method
            text = self.supported_extensions[file_extension](content)
            
            # Generate metadata
            metadata = {
                'file_type': file_extension,
                'file_size': len(content),
                'character_count': len(text),
                'word_count': len(text.split()),
                'filename': filename or (Path(file_path).name if file_path else 'unknown')
            }
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting text from file: {e}")
            raise
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file"""
        if fitz is None:
            raise ImportError("PyMuPDF not available. Please install with: pip install pymupdf")
            
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
                text += "\n\n"  # Separate pages
            
            doc.close()
            return self._clean_extracted_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting from PDF: {e}")
            raise
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(BytesIO(content))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_extracted_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting from DOCX: {e}")
            raise
    
    def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return self._clean_extracted_text(text)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = content.decode('utf-8', errors='replace')
            return self._clean_extracted_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting from TXT: {e}")
            raise
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page breaks and form feeds
        text = text.replace('\f', '\n')
        text = text.replace('\r', '\n')
        
        # Remove common PDF artifacts
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)  # Page numbers
        text = re.sub(r'[\u2000-\u206F\u2E00-\u2E7F]', ' ', text)  # Unicode spaces
        
        return text.strip()
    
    def validate_contract_content(self, text: str) -> Dict[str, any]:
        """
        Validate that the extracted text appears to be a contract
        
        Args:
            text: Extracted text to validate
            
        Returns:
            Dictionary with validation results
        """
        validation_result = {
            'is_likely_contract': False,
            'confidence': 0.0,
            'indicators_found': [],
            'warnings': []
        }
        
        text_lower = text.lower()
        
        # Contract indicators
        contract_indicators = [
            'agreement', 'contract', 'terms and conditions', 'whereas',
            'party', 'parties', 'obligations', 'terms', 'conditions',
            'hereby', 'therefore', 'shall', 'clause', 'section',
            'effective date', 'termination', 'liability', 'indemnify'
        ]
        
        # Legal language patterns
        legal_patterns = [
            r'\b(?:whereas|hereby|therefore|shall|may|must)\b',
            r'\b(?:party|parties)\s+(?:agree|acknowledges?)\b',
            r'\bsection\s+\d+',
            r'\b(?:effective|commencement)\s+date\b',
            r'\btermination\s+(?:clause|provision)\b'
        ]
        
        # Count indicators
        indicators_found = []
        for indicator in contract_indicators:
            if indicator in text_lower:
                indicators_found.append(indicator)
        
        # Count pattern matches
        pattern_matches = 0
        for pattern in legal_patterns:
            if re.search(pattern, text_lower):
                pattern_matches += 1
        
        # Calculate confidence
        indicator_score = len(indicators_found) / len(contract_indicators)
        pattern_score = pattern_matches / len(legal_patterns)
        length_score = min(len(text.split()) / 1000, 1.0)  # Normalize by word count
        
        confidence = (indicator_score * 0.5 + pattern_score * 0.3 + length_score * 0.2)
        
        validation_result['confidence'] = confidence
        validation_result['indicators_found'] = indicators_found
        validation_result['is_likely_contract'] = confidence > 0.3
        
        # Add warnings
        if len(text.split()) < 100:
            validation_result['warnings'].append("Document appears to be very short for a contract")
        
        if confidence < 0.2:
            validation_result['warnings'].append("Document may not be a legal contract")
        
        return validation_result
    
    def segment_contract_sections(self, text: str) -> Dict[str, str]:
        """
        Attempt to identify and segment major contract sections
        
        Args:
            text: Contract text
            
        Returns:
            Dictionary with identified sections
        """
        sections = {
            'preamble': '',
            'definitions': '',
            'terms_and_conditions': '',
            'termination': '',
            'liability': '',
            'miscellaneous': '',
            'signatures': '',
            'unclassified': ''
        }
        
        # Common section headers
        section_patterns = {
            'definitions': r'(?i)(?:definitions?|defined terms)',
            'termination': r'(?i)(?:termination|expiration|end of agreement)',
            'liability': r'(?i)(?:liability|indemnification|damages)',
            'miscellaneous': r'(?i)(?:miscellaneous|general provisions|other terms)',
            'signatures': r'(?i)(?:signatures?|execution|witness)'
        }
        
        lines = text.split('\n')
        current_section = 'unclassified'
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if line is a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_stripped):
                    current_section = section_name
                    section_found = True
                    break
            
            # Add line to current section
            if line_stripped:
                sections[current_section] += line_stripped + '\n'
        
        # If no specific sections found, put everything in terms_and_conditions
        if all(not content.strip() for key, content in sections.items() if key != 'unclassified'):
            sections['terms_and_conditions'] = sections['unclassified']
            sections['unclassified'] = ''
        
        return {k: v.strip() for k, v in sections.items() if v.strip()}
    
    def extract_key_information(self, text: str) -> Dict[str, any]:
        """
        Extract key information from contract text
        
        Args:
            text: Contract text
            
        Returns:
            Dictionary with extracted information
        """
        info = {
            'parties': [],
            'effective_date': None,
            'termination_date': None,
            'governing_law': None,
            'contract_type': None
        }
        
        text_lines = text.split('\n')
        
        # Extract parties (simplified)
        party_patterns = [
            r'(?i)between\s+(.+?)\s+and\s+(.+?)(?:\s+\(|,|\.|$)',
            r'(?i)party\s+(?:a|1|first|one):\s*(.+?)(?:\n|,)',
            r'(?i)party\s+(?:b|2|second|two):\s*(.+?)(?:\n|,)'
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    info['parties'].extend([p.strip() for p in match])
                else:
                    info['parties'].append(match.strip())
        
        # Extract dates
        date_patterns = [
            r'(?i)effective\s+date:?\s*([^\n]+)',
            r'(?i)dated\s+([^\n]+)',
            r'(?i)commencing\s+on\s+([^\n]+)'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match and not info['effective_date']:
                info['effective_date'] = match.group(1).strip()
        
        # Extract governing law
        law_patterns = [
            r'(?i)governed\s+by\s+(?:the\s+)?laws?\s+of\s+([^\n,.]+)',
            r'(?i)jurisdiction\s+of\s+([^\n,.]+)'
        ]
        
        for pattern in law_patterns:
            match = re.search(pattern, text)
            if match:
                info['governing_law'] = match.group(1).strip()
                break
        
        # Determine contract type (basic heuristics)
        type_indicators = {
            'employment': ['employment', 'employee', 'employer', 'salary', 'wage'],
            'service': ['service', 'services', 'provider', 'client', 'deliverables'],
            'sale': ['sale', 'purchase', 'buyer', 'seller', 'goods', 'merchandise'],
            'lease': ['lease', 'rent', 'tenant', 'landlord', 'premises'],
            'license': ['license', 'licensing', 'intellectual property', 'software']
        }
        
        text_lower = text.lower()
        max_score = 0
        detected_type = None
        
        for contract_type, keywords in type_indicators.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > max_score:
                max_score = score
                detected_type = contract_type
        
        info['contract_type'] = detected_type if max_score > 0 else 'unknown'
        
        # Clean up parties list
        info['parties'] = list(set([p for p in info['parties'] if p and len(p) > 2]))[:5]  # Limit to 5 parties
        
        return info
